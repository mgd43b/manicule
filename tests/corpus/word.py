"""Generated `.docx` fixtures: typical, structurally hard, degenerate and hostile.

Generated rather than committed (``docs/parsing.md`` §3.5). The repository stays small, every
fixture's structure is reviewable as code instead of as an opaque binary, and the hostile cases
— a truncated part, a package that is only a zip — exist without ever being stored.

Two choices here are deliberate and worth knowing before editing:

**No heading text repeats within a fixture that the round-trip harness runs over.** A
``heading`` block's text *is* its heading, so two sections called "Overview" make the second
section's resolved text contain the first section's whole block text, and assertion 3 of
``docs/parsing.md`` §3.3 reads that as two locations that cannot be told apart. Slug
de-duplication is exercised instead by headings whose text differs while their slugs collide —
"Roll-back" and "Roll back" — which is the case that actually breaks a fragment.

**Every block's text is unique and not a substring of another section's.** Discrimination
compares text, so a fixture that says "Check the logs" in two sections fails a parser that is
behaving perfectly.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import docx
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.table import Table

__all__ = ["build"]


_FIXED_TIMESTAMP = (2026, 1, 1, 0, 0, 0)
"""One timestamp for every zip member, so the package bytes are the same on every run.

An OOXML file is a zip, and :mod:`zipfile` fills a member's ``date_time`` from the local
clock whenever it is given a name rather than a :class:`zipfile.ZipInfo`. Both python-docx
and python-pptx do exactly that, so without this every fixture differs between two builds and
between two timezones — and a corpus that is not byte-reproducible cannot be compared between
runs, which is the whole basis of asserting anything about it.
"""


def _fixed_timestamps(package: bytes) -> bytes:
    """The same package with one timestamp on every member, in the same member order."""
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(package)) as source, zipfile.ZipFile(out, "w") as target:
        for info in source.infolist():
            entry = zipfile.ZipInfo(info.filename, date_time=_FIXED_TIMESTAMP)
            entry.compress_type = info.compress_type
            entry.external_attr = info.external_attr
            target.writestr(entry, source.read(info.filename))
    return out.getvalue()


def _save_document(package: Document, path: Path) -> None:
    """Write a package to disk with reproducible member timestamps."""
    buffer = io.BytesIO()
    package.save(buffer)
    path.write_bytes(_fixed_timestamps(buffer.getvalue()))


_ASTRAL_HEADING = "配置 𝔘nicode"  # noqa: RUF001 - astral letters are the point of this fixture
"""CJK plus a mathematical-alphanumeric letter: both outside Latin-1, both sluggable."""

_EMOJI_HEADING = "🌏🌍🌎"
"""A heading with nothing sluggable in it, so its fragment is ``None`` and its path is the
only address — the branch that decides whether a repeated path is citable at all."""


def build(dest: Path) -> None:
    """Write this format's fixtures into ``dest``."""
    dest.mkdir(parents=True, exist_ok=True)
    _typical(dest / "word_typical.docx")
    _structurally_hard(dest / "word_structurally_hard.docx")
    _untitled_preamble(dest / "word_untitled_preamble.docx")
    _heading_only(dest / "word_degenerate_heading_only.docx")
    _empty(dest / "word_degenerate_empty.docx")
    (dest / "word_degenerate_zero_bytes.docx").write_bytes(b"")
    _astral(dest / "word_hostile_astral.docx")
    _repeated_heading_path(dest / "word_repeated_heading_path.docx")
    _large(dest / "word-large.docx")
    _truncated(dest / "word_hostile_truncated.docx")
    _plain_zip(dest / "word_hostile_plain_zip.docx")


def _typical(path: Path) -> None:
    document = docx.Document()
    document.core_properties.title = "Deployment Runbook"
    document.add_paragraph("This runbook covers the release train and who to wake.")
    document.add_heading("Deployment", level=1)
    document.add_paragraph("Cut a release branch, then promote it through staging.")
    document.add_paragraph("Announce the window in the release channel.", style="List Bullet")
    document.add_paragraph("Drain the queue before promoting.", style="List Bullet")
    document.add_heading("Rollback", level=2)
    document.add_paragraph("Demote the previous artefact and restore the schema snapshot.")
    table = document.add_table(rows=3, cols=3)
    _fill(
        table,
        [
            ["Step", "Owner", "Timeout"],
            ["Demote artefact", "release engineer", "4 minutes"],
            ["Restore snapshot", "database on-call", "17 minutes"],
        ],
    )
    document.add_heading("Verification", level=1)
    document.add_paragraph("Watch error budget burn for one hour after promotion.")
    _save_document(document, path)


def _structurally_hard(path: Path) -> None:
    document = docx.Document()
    document.core_properties.title = "Platform Handbook"
    _add_deep_list_styles(document)

    document.add_heading("Networking", level=1)
    # A level jump: h1 straight to h3, which must produce a two-element path rather than one
    # padded with an empty string.
    document.add_heading("Peering", level=3)
    document.add_paragraph("Peer sessions are pinned to two transit providers.")

    document.add_heading("Roll-back", level=2)
    document.add_paragraph("Hyphenated section: reverts the network policy generation.")
    document.add_heading("Roll back", level=2)
    document.add_paragraph("Spaced section: reverts the routing table generation.")

    document.add_heading("Nesting", level=1)
    for level in range(1, 6):
        style = "List Bullet" if level == 1 else f"List Bullet {level}"
        document.add_paragraph(f"Depth {level} item, {'sub' * level}ordinate", style=style)

    document.add_heading("Merged", level=1)
    merged = document.add_table(rows=3, cols=4)
    _fill(
        merged,
        [
            ["Quarter", "", "Region", "Owner"],
            ["Q1 plan", "Q1 actual", "EMEA", "Priya"],
            ["Q2 plan", "Q2 actual", "APAC", "Tomás"],
        ],
    )
    # A horizontally merged header cell: it stores its text once and covers two columns, which
    # is what makes repeating a header into a split part non-trivial.
    merged.cell(0, 0).merge(merged.cell(0, 1))

    document.add_heading("Long", level=1)
    lead = document.add_paragraph("The table below is longer than one printed page.")
    # An explicit page break: a lower bound on the page count, and never a pagination — which
    # is why the parser reports no page number for any of this.
    lead.add_run().add_break(WD_BREAK.PAGE)
    long_table = document.add_table(rows=41, cols=3)
    _fill(
        long_table,
        [["Index", "Checkpoint", "Latency"]]
        + [
            [f"{index:03d}", f"checkpoint {index:03d} verified", f"{index * 7} ms"]
            for index in range(1, 41)
        ],
    )
    _save_document(document, path)


def _untitled_preamble(path: Path) -> None:
    """Content before the first heading in a document that declares no title.

    There is nothing to address it with, so it is the fixture behind the 0.05 unlocated budget.
    """
    document = docx.Document()
    document.add_paragraph("Loose note with no heading and no title above it.")
    document.add_heading("Findings", level=1)
    document.add_paragraph("The retry storm began when the health check timeout was lowered.")
    _save_document(document, path)


def _heading_only(path: Path) -> None:
    document = docx.Document()
    document.core_properties.title = "Stub"
    document.add_heading("Placeholder section", level=1)
    _save_document(document, path)


def _empty(path: Path) -> None:
    """A well-formed package with no content. Zero blocks, and that is not a failure."""
    _save_document(docx.Document(), path)


def _astral(path: Path) -> None:
    document = docx.Document()
    document.core_properties.title = "Astral"
    document.add_heading(_ASTRAL_HEADING, level=1)
    # The ambiguous-character rule catches homoglyph typos; these are deliberate, because a
    # citation has to reproduce them exactly.
    document.add_paragraph("Above the basic multilingual plane: 🜁 𝕏 and 𠀋.")  # noqa: RUF001
    table = document.add_table(rows=2, cols=2)
    _fill(table, [["Symbol", "Plane"], ["🌐", "supplementary"]])
    document.add_heading(_EMOJI_HEADING, level=1)
    document.add_paragraph("This section's heading has no sluggable characters at all.")
    _save_document(document, path)


def _repeated_heading_path(path: Path) -> None:
    """Two sections with the identical heading path, and two with an unsluggable one.

    Deliberately outside the six-assertion harness: identical heading text means one section's
    resolved span contains the other's whole heading block, which assertion 3 cannot
    distinguish from a misplaced anchor. The de-duplication and the ambiguity refusal are
    asserted directly in ``tests/parsers/test_word.py`` instead.
    """
    document = docx.Document()
    document.core_properties.title = "Repeats"
    document.add_heading("Service", level=1)
    document.add_heading("Configuration", level=2)
    document.add_paragraph("First configuration section.")
    document.add_heading("Configuration", level=2)
    document.add_paragraph("Second configuration section.")
    document.add_heading("🌏", level=1)
    document.add_paragraph("First unsluggable section.")
    document.add_heading("🌏", level=1)
    document.add_paragraph("Second unsluggable section.")
    _save_document(document, path)


def _large(path: Path) -> None:
    """Bigger than a hand-written fixture, to exercise the streaming path."""
    document = docx.Document()
    document.core_properties.title = "Capacity Review"
    for section in range(1, 9):
        document.add_heading(f"Region {section:02d}", level=1)
        for paragraph in range(1, 26):
            document.add_paragraph(
                f"Region {section:02d} note {paragraph:02d}: sustained throughput held at "
                f"{section * 1000 + paragraph} requests per second with no queue growth, and "
                f"the p99 stayed inside the objective for the whole window."
            )
    _save_document(document, path)


def _add_deep_list_styles(document: Document) -> None:
    """Add the level 4 and 5 list styles Word's default template stops short of.

    Word ships List Bullet 1-3 only, so a five-deep list needs two more styles, which is what a
    house template does. The parser reads the depth off the style name rather than off
    indentation, because indentation is a rendering property and a template can set it to
    anything.

    ``add_style`` needs a suppression: python-docx leaves its parameters unannotated, so the
    checker reports the method's type as partially unknown. That is a statement about upstream,
    not about this call, and the arguments below are exactly what its docstring asks for.
    """
    kind = WD_STYLE_TYPE.PARAGRAPH
    for level in (4, 5):
        name = f"List Bullet {level}"
        document.styles.add_style(name, kind)  # pyright: ignore[reportUnknownMemberType] - see above


def _fill(table: Table, rows: list[list[str]]) -> None:
    """Write a grid into a table, cell by cell."""
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value


def _truncated(path: Path) -> None:
    """A package whose ``word/document.xml`` stops mid-attribute. The parser must decline."""
    buffer = io.BytesIO()
    document = docx.Document()
    document.add_heading("Intact heading", level=1)
    document.add_paragraph("Intact paragraph that the truncated copy loses.")
    document.save(buffer)
    path.write_bytes(_truncate_member(buffer.getvalue(), "word/document.xml"))


def _plain_zip(path: Path) -> None:
    """A zip that is not an OOXML package at all, under a `.docx` name."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            zipfile.ZipInfo("notes.txt", date_time=_FIXED_TIMESTAMP),
            "not an office document",
        )
    path.write_bytes(buffer.getvalue())


def _truncate_member(package: bytes, member: str) -> bytes:
    """Rewrite a zip with one member cut in half, leaving the container itself valid."""
    source = zipfile.ZipFile(io.BytesIO(package))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as archive:
        for info in source.infolist():
            blob = source.read(info.filename)
            entry = zipfile.ZipInfo(info.filename, date_time=_FIXED_TIMESTAMP)
            entry.compress_type = info.compress_type
            archive.writestr(entry, blob[: len(blob) // 2] if info.filename == member else blob)
    return out.getvalue()
